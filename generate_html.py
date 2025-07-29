import math

from docx import Document
from docx.shared import Length
import os
import sys
import re


def getOutlineLevel(inputXml):
    """
    功能 从xml字段中提取出<w:outlineLvl w:val="number"/>中的数字number
    参数 inputXml
    返回 number
    """
    start_index = inputXml.find('<w:outlineLvl')
    end_index = inputXml.find('>', start_index)
    number = inputXml[start_index:end_index + 1]
    number = re.search("\d+", number).group()
    return int(number) + 1


def isTitle(paragraph):
    """
    功能 判断该段落是否设置了大纲等级
    参数 paragraph:段落
    返回 None:普通正文，没有大纲级别 0:一级标题 1:二级标题 2:三级标题
    """
    # 如果是空行，直接返回None
    if paragraph.text.strip() == '':
        return None

    # 如果该段落是直接在段落里设置大纲级别的，根据xml判断大纲级别
    paragraphXml = paragraph._p.xml
    if paragraphXml.find('<w:outlineLvl') >= 0:
        return getOutlineLevel(paragraphXml)
    # 如果该段落是通过样式设置大纲级别的，逐级检索样式及其父样式，判断大纲级别
    # 之所以要检查父样式。原因是如果a样式是大纲1，那么如果a样式改了下字体字号，然后就生成了样式b,那么此时样式b就没有<w:outlineLvl字段。
    targetStyle = paragraph.style
    while targetStyle is not None:
        # 如果在该级style中找到了大纲级别，返回
        if targetStyle.element.xml.find('<w:outlineLvl') >= 0:
            return getOutlineLevel(targetStyle.element.xml)
        else:
            targetStyle = targetStyle.base_style
    # 如果在段落、样式里都没有找到大纲级别，返回None
    return None

# 弃用
def get_heading_level(paragraph):
    # 通过大纲级别判断
    # outline_level 不存在这个成员说是。所以这个方法作废
    outline_level = paragraph.paragraph_format.outline_level
    if outline_level is not None:
        level_name = outline_level.name
        if level_name.startswith('LEVEL_'):
            try:
                return int(level_name.split('_')[1])
            except (IndexError, ValueError):
                pass

    # 通过样式名称判断（支持中英文）
    style_name = paragraph.style.name
    match = re.match(r'^(?:Heading|标题)\s*(\d+)$', style_name, re.IGNORECASE)
    if match:
        try:
            return int(match.group(1))
        except (IndexError, ValueError):
            pass

    return None

# 判断标题几的示例用法
def analyze_headings(file_path):
    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"无法打开文件: {e}")
        return

    heading_count = 0

    for idx, paragraph in enumerate(doc.paragraphs):
        level = isTitle(paragraph)
        if level is not None:
            heading_count += 1
            text = paragraph.text.strip()[:50]  # 截取前50个字符防止过长
            print(f"[段落 {idx + 1}] 标题级别 {level} ({text}...)")

    if heading_count == 0:
        print("未发现任何标题内容")
    else:
        print(f"\n共找到 {heading_count} 个标题")




def get_effective_indent_pt(paragraph, indent_attr):
    """
    获取段落有效的缩进值（以磅为单位）
    :param paragraph: 段落对象
    :param indent_attr: 缩进属性名称（'left_indent', 'right_indent', 'first_line_indent'）
    :return: 缩进值（磅）
    """
    try:
        # 检查直接格式
        direct_value = getattr(paragraph.paragraph_format, indent_attr)
        if direct_value is not None and isinstance(direct_value, Length):
            return direct_value.pt
    except AttributeError:
        pass

    try:
        # 检查样式继承链
        style = paragraph.style
        current_style = style
        while current_style is not None:
            style_pf = current_style.paragraph_format
            style_value = getattr(style_pf, indent_attr)
            if style_value is not None and isinstance(style_value, Length):
                return style_value.pt
            current_style = current_style.base_style
    except AttributeError:
        pass

    return 0.0  # 默认值

# 判断缩进的示例用法
def print_paragraph_indents(docx_path):
    doc = Document(docx_path)

    for para_idx, paragraph in enumerate(doc.paragraphs, 1):
        # 获取三种缩进类型
        indent_types = [
            ('左缩进', 'left_indent'),
            ('首行缩进', 'first_line_indent'),
            ('右缩进', 'right_indent')
        ]

        indent_values = []
        for label, attr in indent_types:
            value = get_effective_indent_pt(paragraph, attr)
            indent_values.append(f"{label}: {value:.2f}磅")

        # 构建输出信息
        output = [
            f"段落 {para_idx}",
            f"文本内容: {paragraph.text[:40] + '...' if len(paragraph.text) > 40 else paragraph.text}"
        ]
        output.extend(indent_values)

        print("\n".join(output))
        print("-" * 50)


if __name__ == "__main__":

    # 获取当前工作目录
    current_dir = os.getcwd()

    # 遍历当前目录及所有子目录
    for root, dirs, files in os.walk(current_dir):
        for file in files:
            # 检查文件是否以.docx结尾（不区分大小写）
            if file.lower().endswith('.docx'):
                if file.lower()[0:2] != "~$":
                    # 拼接文件的绝对路径
                    file_path = os.path.join(root, file)
                    html_file_path = os.path.join(root,file[:-5]+".html")
                    # 打印绝对路径
                    print(file_path)
                    print(html_file_path)

                    # 检查文件是否存在
                    if os.path.exists(html_file_path):
                        # 如果存在，则删除文件
                        os.remove(html_file_path)

                    html_file = open(html_file_path,'w+',encoding='utf-8')
                    html_file.write("""<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>""")
                    html_file.write(file[:-5])
                    html_file.write("""</title>
    <style>
        * {
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }
        
        body {
            font-family: 'Segoe UI', 'Microsoft YaHei', sans-serif;
            line-height: 1.6;
            color: #333;
            background: linear-gradient(135deg, #f5f7fa 0%, #c3cfe2 100%);
            min-height: 100vh;
            padding: 20px;
        }
        
        .container {
            display: flex;
            max-width: 1400px;
            margin: 0 auto;
            gap: 30px;
        }
        
        header {
            text-align: center;
            margin-bottom: 40px;
            padding: 20px;
            background: rgba(255, 255, 255, 0.8);
            border-radius: 12px;
            box-shadow: 0 4px 15px rgba(0, 0, 0, 0.1);
        }

        blockquote {
            padding-left: 45px;
        }
        
        h1 {
            font-size: 2.8rem;
            margin-bottom: 10px;
            color: #2c3e50;
            background: linear-gradient(90deg, #3498db, #8e44ad);
            -webkit-background-clip: text;
            background-clip: text;
            color: transparent;
        }
        h2 {
            font-size: 2rem;
            margin-bottom: 20px;
            padding-bottom: 10px;
            border-bottom: 2px solid #3498db;
            color: #2c3e50;
        }
        
        h3 {
            font-size: 1.6rem;
            margin: 25px 0 15px;
            color: #2050aa;
        }
        
        h4 {
            font-size: 1.3rem;
            margin: 20px 0 12px;
            color: #9b59b6;
        }
        
        h5 {
            font-size: 1.1rem;
            margin: 18px 0 10px;
            color: #e67e22;
        }
        
        h6 {
            font-size: 1rem;
            margin: 16px 0 8px;
            color: #e74c3c;
        }
        
        p {
            margin-bottom: 15px;
            font-size: 1.05rem;
            color: #555;
            line-height: 1.7;
        }


        .subtitle {
            font-size: 1.2rem;
            color: #7f8c8d;
            max-width: 800px;
            margin: 0 auto;
        }
        
        /* 侧边栏样式 */
        .sidebar {
            flex: 0 0 280px;
            background: rgba(255, 255, 255, 0.95);
            border-radius: 12px;
            padding: 25px 20px;
            box-shadow: 0 8px 30px rgba(0, 0, 0, 0.12);
            height: fit-content;
            position: sticky;
            top: 20px;
            max-height: 90vh;
            overflow-y: auto;
        }
        
        .sidebar-title {
            font-size: 1.6rem;
            margin-bottom: 20px;
            padding-bottom: 10px;
            border-bottom: 2px solid #3498db;
            color: #2c3e50;
            display: flex;
            align-items: center;
            gap: 10px;
        }
        
        .sidebar-title::before {
            content: "📋";
        }
        
        #toc {
            list-style: none;
            
            padding-left: 0;
        }
        #toc ul {
            list-style: none; /* 必须添加 */
        }
        
        
        #toc li {
            margin-bottom: 8px;
            transition: all 0.3s ease;
        }
        
        #toc a {
            text-decoration: none;
            color: #34495e;
            display: block;
            padding: 8px 15px;
            border-radius: 8px;
            transition: all 0.2s ease;
            font-weight: 500;
        }
        
        #toc a:hover {
            background: #e3f2fd;
            color: #2980b9;
            transform: translateX(5px);
        }
        #toc .h1 a::before { content: "■ "; color: #3498db; }
        #toc .h2 a::before { content: "► "; color: #9b59b6; }
        #toc .h3 a::before { content: "▸ "; color: #2ecc71; }
        #toc .h1 a {
            font-size: 1.2rem;
            font-weight: 600;
            color: #2980b9;
            border-left: 4px solid #3498db;

        }
        
        #toc .h2 a {
            font-size: 1.1rem;
            padding-left: 30px;
            font-weight: normal;
            color: #34495e;
            font-family: 'Segoe UI', 'Microsoft YaHei', sans-serif;
            border-left: 3px solid #9b59b6;

        }

        #toc .h3 a {
            font-size: 1rem;
            padding-left: 45px;
            color: #34495e;
            font-family: 'Segoe UI', 'Microsoft YaHei', sans-serif;
            border-left: 2px solid #2ecc71;
        }
        
        #toc .h4 a {
            font-size: 0.95rem;
            padding-left: 60px;
            color: #34495e;
            font-family: 'Segoe UI', 'Microsoft YaHei', sans-serif;
            border-left: 2px solid #f39c12;
        }
        
        #toc .h5 a {
            font-size: 0.9rem;
            padding-left: 75px;
            color: #34495e;
            font-family: 'Segoe UI', 'Microsoft YaHei', sans-serif;
            border-left: 2px solid #e74c3c;
        }
        #toc .h6 a {
            font-size: 0.85rem;
            padding-left: 90px;
            color: #34495e;
            font-family: 'Segoe UI', 'Microsoft YaHei', sans-serif;
            border-left: 2px solid #e74c3c;
        }

        
        /* 主内容区样式 */
        .content {
            flex: 1;
            background: rgba(255, 255, 255, 0.95);
            border-radius: 12px;
            padding: 40px;
            box-shadow: 0 8px 30px rgba(0, 0, 0, 0.12);
        }
        
        .content-section {
            margin-bottom: 50px;
        }
        
        .content-section h2 {
            font-size: 2rem;
            margin-bottom: 20px;
            padding-bottom: 10px;
            border-bottom: 2px solid #3498db;
            color: #2c3e50;
        }
        
        .content-section h3 {
            font-size: 1.6rem;
            margin: 25px 0 15px;
            color: #2980b9;
        }
        
        .content-section h4 {
            font-size: 1.3rem;
            margin: 20px 0 12px;
            color: #9b59b6;
        }
        
        .content-section h5 {
            font-size: 1.1rem;
            margin: 18px 0 10px;
            color: #e67e22;
        }
        
        .content-section h6 {
            font-size: 1rem;
            margin: 16px 0 8px;
            color: #e74c3c;
        }
        
        .content-section p {
            margin-bottom: 15px;
            font-size: 1.05rem;
            color: #555;
            line-height: 1.7;
        }
        
        .code-block {
            background: #2c3e50;
            color: #ecf0f1;
            padding: 20px;
            border-radius: 8px;
            margin: 20px 0;
            font-family: 'Courier New', monospace;
            overflow-x: auto;
        }
        
        .highlight {
            background-color: rgba(255, 255, 0, 0.2);
            padding: 2px 5px;
            border-radius: 4px;
        }
        
        /* 响应式设计 */
        @media (max-width: 900px) {
            .container {
                flex-direction: column;
            }
            
            .sidebar {
                position: static;
                width: 100%;
            }
        }
        
        .back-to-top {
            position: fixed;
            bottom: 30px;
            right: 30px;
            background: #3498db;
            color: white;
            width: 50px;
            height: 50px;
            border-radius: 50%;
            display: flex;
            align-items: center;
            justify-content: center;
            font-size: 1.5rem;
            cursor: pointer;
            box-shadow: 0 4px 10px rgba(0, 0, 0, 0.2);
            transition: all 0.3s ease;
            opacity: 0;
            transform: translateY(20px);
            z-index: 100;
        }
        
        .back-to-top.show {
            opacity: 1;
            transform: translateY(0);
        }
        
        .back-to-top:hover {
            background: #2980b9;
            transform: scale(1.1);
        }
        
        footer {
            text-align: center;
            margin-top: 40px;
            padding: 20px;
            color: #7f8c8d;
            font-size: 0.9rem;
        }
    </style>
</head>
<body>
    <div class="container">
        <aside class="sidebar">
            <h2 class="sidebar-title">文档目录</h2>
            <ul id="toc"></ul>
        </aside>
        
        <main class="content">""")


                    doc = Document(file_path)

                    for para_idx, paragraph in enumerate(doc.paragraphs, 1):

                        if len(paragraph.text) == 0:
                            html_file.write("<br><br>\n")

                        else:
                            level = isTitle(paragraph)
                            if level is not None:
                                string_start = '<h' + str(level) + '>'
                                string_end = '</h' + str(level) + '>'

                                string = string_start + paragraph.text + string_end + '\n'
                                html_file.write(string)

                            else:
                                # 获取三种缩进类型
                                indent_types = [
                                    ('左缩进', 'left_indent'),
                                    ('首行缩进', 'first_line_indent'),
                                    ('右缩进', 'right_indent')
                                ]

                                indent_values = []
                                value = get_effective_indent_pt(paragraph, indent_types[0][1])

                                n = math.floor(value/20)
                                if n > 0:
                                    string_start = '<blockquote>'*n
                                    string_end = '</blockquote>'*n
                                    string = string_start + paragraph.text + string_end
                                else:
                                    string = '<p>' + paragraph.text + '</p>'

                                html_file.write(string+'\n')

                    html_file.write("""        </main>
    </div>
    
    <div class="back-to-top" id="backToTop">↑</div>
    
    <footer>
        <p>已经到底啦</p>
    </footer>
    
    <script>
        document.addEventListener('DOMContentLoaded', function() {
            // 生成目录
            generateTOC();
            
            // 设置滚动高亮
            window.addEventListener('scroll', highlightActiveHeading);
            
            // 初始化返回顶部按钮
            initBackToTop();
            
            // 添加平滑滚动
            addSmoothScrolling();

            
        });
        
        function generateTOC() {
            const tocContainer = document.getElementById('toc');
            const headings = document.querySelectorAll('.content h1, .content h2, .content h3, .content h4, .content h5, .content h6');
            
            // 用于跟踪当前层级
            let currentLevel = 0;
            // 目录栈，用于处理嵌套层级
            const tocStack = [tocContainer];
            
            // 清空目录容器
            tocContainer.innerHTML = '';
            
            // 遍历所有标题元素
            headings.forEach((heading, index) => {
                // 确保标题有ID
                if (!heading.id) {
                    heading.id = 'heading-' + index;
                }
                
                // 获取标题级别（h1->1, h2->2, ...）
                const level = parseInt(heading.tagName.substring(1));
                
                // 创建列表项
                const li = document.createElement('li');
                li.className = `h${level}`;
                
                // 创建链接
                const a = document.createElement('a');
                a.href = `#${heading.id}`;
                a.textContent = heading.textContent;
                a.dataset.level = level;
                
                li.appendChild(a);
                
                // 处理层级关系
                if (level > currentLevel) {
                    // 进入更深层级 - 创建新的ul
                    const ul = document.createElement('ul');
                    li.appendChild(ul);
                    tocStack[tocStack.length - 1].appendChild(li);
                    tocStack.push(ul);
                } else if (level < currentLevel) {
                    // 返回上层级 - 弹出栈
                    const popCount = currentLevel - level;
                    for (let i = 0; i < popCount; i++) {
                        if (tocStack.length > 1) tocStack.pop();
                    }
                    tocStack[tocStack.length - 1].appendChild(li);
                } else {
                    // 同级 - 直接添加到当前ul
                    tocStack[tocStack.length - 1].appendChild(li);
                }
                
                // 更新当前层级
                currentLevel = level;
            });
        }
        
        function highlightActiveHeading() {
            const headings = document.querySelectorAll('.content h1, .content h2, .content h3, .content h4, .content h5, .content h6');
            const tocLinks = document.querySelectorAll('#toc a');
            
            // 移除所有活动类
            tocLinks.forEach(link => link.classList.remove('active'));
            
            // 找到当前视口最顶部的标题
            let activeHeading = null;
            let minDistance = Infinity;
            
            headings.forEach(heading => {
                const rect = heading.getBoundingClientRect();
                // 计算标题顶部与视口顶部的距离
                const distance = rect.top;
                
                // 如果标题在视口顶部附近，且距离更小
                if (distance >= 0 && distance < minDistance) {
                    minDistance = distance;
                    activeHeading = heading;
                }
            });
            
            // 高亮对应的目录项
            if (activeHeading) {
                const activeLink = document.querySelector(`#toc a[href="#${activeHeading.id}"]`);
                if (activeLink) {
                    activeLink.classList.add('active');
                    
                    // 滚动到活动链接位置
                    activeLink.scrollIntoView({ behavior: 'smooth', block: 'nearest', inline: 'start' });
                }
            }
            
            // 处理返回顶部按钮
            const backToTop = document.getElementById('backToTop');
            if (window.scrollY > 300) {
                backToTop.classList.add('show');
            } else {
                backToTop.classList.remove('show');
            }
        }
        
        function initBackToTop() {
            const backToTop = document.getElementById('backToTop');
            backToTop.addEventListener('click', () => {
                window.scrollTo({
                    top: 0,
                    behavior: 'smooth'
                });
            });
        }
        
        function addSmoothScrolling() {
            // 为目录链接添加平滑滚动
            document.querySelectorAll('#toc a').forEach(link => {
                link.addEventListener('click', function(e) {
                    e.preventDefault();
                    const targetId = this.getAttribute('href');
                    const targetElement = document.querySelector(targetId);
                    
                    if (targetElement) {
                        window.scrollTo({
                            top: targetElement.offsetTop - 30,
                            behavior: 'smooth'
                        });
                        
                        // 更新URL hash
                        history.replaceState(null, null, targetId);
                    }
                });
            });
        }
        
    </script>
</body>
</html>""")
                    html_file.close()

                    # print_paragraph_indents(file_path)
                    # analyze_headings(file_path)







    # print_paragraph_indents(sys.argv[1])